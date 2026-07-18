# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "pymupdf",
#   "python-docx",
#   "fastembed",
#   "duckdb",
#   "numpy",
# ]
# ///
"""Zolam v3 flat-file ingest/update/query pipeline.

Invoked by the zolam Go CLI via `uv run ingest.py`. Owns extraction,
chunking, embedding, and all reads/writes of the per-project index
(index.duckdb or index.jsonl). Human-readable progress goes to stderr;
the final line on stdout is always a single JSON object.

    uv run ingest.py --mode ingest --project-dir <dir> --backend duckdb \
        --files <path> [<path> ...] --removed <path> [<path> ...]

    uv run ingest.py --mode query --project-dir <dir> --backend duckdb \
        --query "<text>" --top-k 5 [--keyword]
"""

import argparse
import functools
import hashlib
import json
import os
import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

import duckdb
import numpy as np

DEFAULT_MODEL = "BAAI/bge-small-en-v1.5"
DEFAULT_DIMS = 384
SCRIPT_VERSION = "3"

CHUNK_SIZE = 2000
OVERLAP = int(CHUNK_SIZE * 0.15)  # ~300 chars, ~15%

BINARY_EXTENSIONS = {".pdf", ".docx"}


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def _split_paragraphs(text: str) -> list[str]:
    return re.split(r"\n\s*\n", text)


def _is_heading(paragraph: str) -> bool:
    return paragraph.lstrip().startswith("#")


def _hard_split(s: str, size: int = CHUNK_SIZE, overlap: int = OVERLAP) -> list[str]:
    """Split a single oversized unit into fixed-size, overlapping pieces."""
    if len(s) <= size:
        return [s]
    chunks = []
    start = 0
    n = len(s)
    while start < n:
        end = start + size
        chunks.append(s[start:end])
        if end >= n:
            break
        start = end - overlap
    return chunks


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = OVERLAP) -> list[str]:
    """Split text into ~size-char chunks with ~overlap-char continuity.

    Prefers markdown heading boundaries, then paragraph (blank-line)
    boundaries, then hard-splits any single paragraph too large to fit.
    """
    text = text.strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]

    paragraphs = [p for p in _split_paragraphs(text) if p.strip()]
    chunks: list[str] = []
    buffer = ""

    def start_buffer(seed_para: str) -> str:
        tail = chunks[-1][-overlap:] if chunks else ""
        return (tail + "\n\n" + seed_para) if tail else seed_para

    for para in paragraphs:
        if len(para) > size:
            if buffer.strip():
                chunks.append(buffer.strip())
                buffer = ""
            chunks.extend(_hard_split(para, size, overlap))
            continue

        if _is_heading(para) and buffer.strip():
            chunks.append(buffer.strip())
            buffer = start_buffer(para)
            continue

        candidate = (buffer + "\n\n" + para) if buffer else para
        if len(candidate) > size:
            chunks.append(buffer.strip())
            buffer = start_buffer(para)
        else:
            buffer = candidate

    if buffer.strip():
        chunks.append(buffer.strip())

    return chunks if chunks else [text]


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------

def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def read_plain(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


@functools.lru_cache(maxsize=1)
def _find_tessdata_dir() -> str | None:
    """Locate a Tesseract `tessdata` folder (the language files PyMuPDF's
    OCR needs) since Tesseract's own installers don't set TESSDATA_PREFIX.
    Only the language data is required, not a full Tesseract install.
    Checked in order: TESSDATA_PREFIX, then layouts relative to a
    `tesseract` binary on PATH (following scoop's shim indirection on
    Windows), then well-known per-OS default paths."""
    env = os.environ.get("TESSDATA_PREFIX")
    if env and Path(env).is_dir():
        return env

    candidates: list[Path] = []
    tess_bin = shutil.which("tesseract")
    if tess_bin:
        bin_path = Path(tess_bin)
        # scoop shims are stub executables; the real install path lives in
        # a sibling ".shim" file rather than being a filesystem symlink.
        shim_file = bin_path.with_suffix(".shim")
        if shim_file.exists():
            for line in shim_file.read_text(encoding="utf-8", errors="ignore").splitlines():
                if line.strip().lower().startswith("path"):
                    real = line.split("=", 1)[-1].strip().strip('"')
                    if real:
                        bin_path = Path(real)
                    break
        bin_dir = bin_path.resolve().parent
        candidates += [
            bin_dir / "tessdata",
            bin_dir.parent / "tessdata",
            bin_dir.parent / "share" / "tessdata",
            bin_dir.parent / "share" / "tesseract-ocr" / "tessdata",
        ]
        candidates += sorted(bin_dir.parent.glob("share/tesseract-ocr/*/tessdata"))

    scoop_dir = os.environ.get("SCOOP") or str(Path.home() / "scoop")
    candidates += [
        Path(scoop_dir) / "apps" / "tesseract" / "current" / "tessdata",
        Path(os.environ.get("ProgramFiles", "C:/Program Files")) / "Tesseract-OCR" / "tessdata",
        Path("/usr/share/tesseract-ocr/5/tessdata"),
        Path("/usr/share/tesseract-ocr/4.00/tessdata"),
        Path("/usr/share/tessdata"),
        Path("/opt/homebrew/share/tessdata"),
        Path("/usr/local/share/tessdata"),
    ]

    for candidate in candidates:
        if candidate.is_dir() and any(candidate.glob("*.traineddata")):
            return str(candidate)
    return None


def extract_pdf_pages(path: Path) -> list[str]:
    """Extract each page's text layer, falling back to OCR (via PyMuPDF's
    bundled Tesseract engine) for pages with no embedded text, e.g. scanned
    PDFs. OCR needs a `tessdata` language-file folder locatable via
    TESSDATA_PREFIX or a Tesseract install (see _find_tessdata_dir); if
    unavailable the page is left as its (empty) text-layer result rather
    than failing the whole file."""
    import fitz  # pymupdf

    doc = fitz.open(str(path))
    try:
        pages = []
        tessdata = _find_tessdata_dir()
        for page in doc:
            text = page.get_text()
            if not text.strip():
                try:
                    ocr_kwargs = {"language": "eng"}
                    if tessdata:
                        ocr_kwargs["tessdata"] = tessdata
                    text = page.get_text(textpage=page.get_textpage_ocr(**ocr_kwargs))
                except Exception as ocr_err:
                    hint = "" if tessdata else " (install Tesseract, or set TESSDATA_PREFIX to a tessdata folder — see README)"
                    print(f"  OCR failed {path} page {page.number + 1}: {ocr_err}{hint}", file=sys.stderr)
            pages.append(text)
        return pages
    finally:
        doc.close()


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _sidecar_path(project_dir: Path, source_path: Path) -> Path:
    return project_dir / "extracted" / f"{source_path.name}.md"


def _front_matter(source_path: Path) -> str:
    return (
        "---\n"
        f"source: {source_path}\n"
        f"extracted: {datetime.now(timezone.utc).isoformat()}\n"
        f"sha256: {sha256_file(source_path)}\n"
        "---\n"
    )


def write_pdf_sidecar(project_dir: Path, source_path: Path, pages: list[str]) -> str:
    sidecar = _sidecar_path(project_dir, source_path)
    sidecar.parent.mkdir(parents=True, exist_ok=True)

    parts = [_front_matter(source_path), f"# {source_path.name}\n"]
    for i, text in enumerate(pages, start=1):
        parts.append(f"\n## Page {i}\n")
        parts.append(text.strip() if text.strip() else "*(no extractable text)*")
    sidecar.write_text("\n".join(parts), encoding="utf-8")
    return str(sidecar.relative_to(project_dir))


def write_text_sidecar(project_dir: Path, source_path: Path, text: str) -> str:
    sidecar = _sidecar_path(project_dir, source_path)
    sidecar.parent.mkdir(parents=True, exist_ok=True)
    body = f"# {source_path.name}\n\n{text}"
    sidecar.write_text(_front_matter(source_path) + body, encoding="utf-8")
    return str(sidecar.relative_to(project_dir))


def remove_sidecar(project_dir: Path, source_path: Path) -> None:
    if source_path.suffix.lower() in BINARY_EXTENSIONS:
        _sidecar_path(project_dir, source_path).unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# Embedding
# ---------------------------------------------------------------------------

class Embedder:
    """Lazily loads the fastembed model on first use."""

    def __init__(self, model_name: str):
        self.model_name = model_name
        self._model = None

    def embed(self, texts: list[str]) -> list[list[float]]:
        if self._model is None:
            from fastembed import TextEmbedding

            print(f"Loading embedding model {self.model_name}...", file=sys.stderr)
            self._model = TextEmbedding(model_name=self.model_name)
        return [v.tolist() for v in self._model.embed(texts)]


# ---------------------------------------------------------------------------
# Backends
# ---------------------------------------------------------------------------

class DuckDBBackend:
    def __init__(self, path: str, model: str, dims: int):
        self.dims = dims
        self.con = duckdb.connect(path)
        self.con.execute("CREATE TABLE IF NOT EXISTS meta (key VARCHAR PRIMARY KEY, value VARCHAR)")
        self.con.execute(
            f"CREATE TABLE IF NOT EXISTS chunks ("
            f"path VARCHAR, chunk_num INTEGER, page INTEGER, text VARCHAR, "
            f"embedding FLOAT[{dims}])"
        )
        for key, value in (("model", model), ("dims", str(dims)), ("zolam_version", SCRIPT_VERSION)):
            self.con.execute(
                "INSERT INTO meta VALUES (?, ?) ON CONFLICT (key) DO UPDATE SET value = excluded.value",
                [key, value],
            )

    def delete_paths(self, paths: list[str]) -> None:
        if not paths:
            return
        self.con.executemany("DELETE FROM chunks WHERE path = ?", [[p] for p in paths])

    def insert_chunks(self, records: list[dict]) -> None:
        if not records:
            return
        self.con.executemany(
            "INSERT INTO chunks VALUES (?, ?, ?, ?, ?)",
            [[r["path"], r["chunk"], r["page"], r["text"], r["embedding"]] for r in records],
        )

    def search(self, query_embedding: list[float], top_k: int) -> list[dict]:
        rows = self.con.execute(
            f"SELECT path, chunk_num, page, text, "
            f"array_cosine_similarity(embedding, ?::FLOAT[{self.dims}]) AS score "
            f"FROM chunks ORDER BY score DESC LIMIT ?",
            [query_embedding, top_k],
        ).fetchall()
        return [
            {"path": r[0], "chunk": r[1], "page": r[2], "text": r[3], "score": r[4]}
            for r in rows
        ]

    def keyword_search(self, term: str, top_k: int) -> list[dict]:
        rows = self.con.execute(
            "SELECT path, chunk_num, page, text FROM chunks WHERE text ILIKE ? LIMIT ?",
            [f"%{term}%", top_k],
        ).fetchall()
        return [
            {"path": r[0], "chunk": r[1], "page": r[2], "text": r[3], "score": None}
            for r in rows
        ]

    def close(self) -> None:
        self.con.close()


class JsonlBackend:
    def __init__(self, path: str, model: str, dims: int):
        self.path = path
        self.model = model
        self.dims = dims
        self.records: list[dict] = []
        p = Path(path)
        if p.exists():
            with p.open(encoding="utf-8") as f:
                lines = [line.strip() for line in f if line.strip()]
            for line in lines[1:]:  # skip the _meta header line
                self.records.append(json.loads(line))

    def delete_paths(self, paths: list[str]) -> None:
        if not paths:
            return
        pathset = set(paths)
        self.records = [r for r in self.records if r["path"] not in pathset]

    def insert_chunks(self, records: list[dict]) -> None:
        self.records.extend(records)

    def search(self, query_embedding: list[float], top_k: int) -> list[dict]:
        if not self.records:
            return []
        mat = np.array([r["embedding"] for r in self.records], dtype=np.float32)
        q = np.array(query_embedding, dtype=np.float32)
        q_norm = q / (np.linalg.norm(q) + 1e-10)
        mat_norm = mat / (np.linalg.norm(mat, axis=1, keepdims=True) + 1e-10)
        scores = mat_norm @ q_norm
        order = np.argsort(-scores)[:top_k]
        return [
            {
                "path": self.records[i]["path"],
                "chunk": self.records[i]["chunk"],
                "page": self.records[i].get("page"),
                "text": self.records[i]["text"],
                "score": float(scores[i]),
            }
            for i in order
        ]

    def keyword_search(self, term: str, top_k: int) -> list[dict]:
        term_lower = term.lower()
        matches = [r for r in self.records if term_lower in r["text"].lower()][:top_k]
        return [
            {"path": r["path"], "chunk": r["chunk"], "page": r.get("page"), "text": r["text"], "score": None}
            for r in matches
        ]

    def close(self) -> None:
        tmp = self.path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            f.write(json.dumps({"_meta": True, "version": 3, "model": self.model, "dims": self.dims}) + "\n")
            for r in self.records:
                f.write(json.dumps(r) + "\n")
        Path(tmp).replace(self.path)


def open_backend(project_dir: Path, backend_name: str, model: str, dims: int, require_existing: bool = False):
    if backend_name == "duckdb":
        path = str(project_dir / "index.duckdb")
        if require_existing and not Path(path).exists():
            raise RuntimeError(f"no duckdb index at {path}; run 'zolam ingest' first")
        return DuckDBBackend(path, model, dims)
    if backend_name == "jsonl":
        path = str(project_dir / "index.jsonl")
        if require_existing and not Path(path).exists():
            raise RuntimeError(f"no jsonl index at {path}; run 'zolam ingest' first")
        return JsonlBackend(path, model, dims)
    raise RuntimeError(f"unsupported backend {backend_name!r} (expected duckdb or jsonl)")


# ---------------------------------------------------------------------------
# Ingest / update
# ---------------------------------------------------------------------------

def process_file(path: Path, project_dir: Path, embedder: Embedder) -> list[dict]:
    """Extract, chunk, and embed a single file. Returns chunk records
    (without a backend-assigned path prefix change) ready for insertion."""
    ext = path.suffix.lower()

    if ext == ".pdf":
        pages = extract_pdf_pages(path)
        write_pdf_sidecar(project_dir, path, pages)
        units: list[tuple[int | None, str]] = []
        for page_num, page_text in enumerate(pages, start=1):
            for chunk in chunk_text(page_text):
                units.append((page_num, chunk))
    elif ext == ".docx":
        text = extract_docx(path)
        write_text_sidecar(project_dir, path, text)
        units = [(None, c) for c in chunk_text(text)]
    else:
        text = read_plain(path)
        units = [(None, c) for c in chunk_text(text)]

    if not units:
        return []

    texts = [u[1] for u in units]
    embeddings = embedder.embed(texts)
    records = []
    for i, ((page, text), embedding) in enumerate(zip(units, embeddings)):
        records.append({"path": str(path), "chunk": i, "page": page, "text": text, "embedding": embedding})
    return records


def run_ingest(args: argparse.Namespace) -> None:
    project_dir = Path(args.project_dir)
    project_dir.mkdir(parents=True, exist_ok=True)

    backend = open_backend(project_dir, args.backend, args.embedding_model, args.embedding_dims)
    embedder = Embedder(args.embedding_model)

    removed = [str(Path(p)) for p in (args.removed or [])]
    files = [str(Path(p)) for p in (args.files or [])]

    backend.delete_paths(removed + files)
    for p in removed:
        remove_sidecar(project_dir, Path(p))

    errors: list[str] = []
    chunks_written = 0
    files_processed = 0

    for i, file_str in enumerate(files, start=1):
        path = Path(file_str)
        print(f"[{i}/{len(files)}] {path.name}", file=sys.stderr)
        try:
            records = process_file(path, project_dir, embedder)
            if records:
                backend.insert_chunks(records)
                chunks_written += len(records)
            files_processed += 1
        except Exception as e:  # noqa: BLE001 - reported to caller, not fatal
            errors.append(f"{path}: {e}")
            print(f"  ERROR: {e}", file=sys.stderr)

    backend.close()

    result = {
        "files_processed": files_processed,
        "files_errored": len(errors),
        "files_removed": len(removed),
        "chunks_written": chunks_written,
        "errors": errors,
    }
    print(json.dumps(result))


def run_query(args: argparse.Namespace) -> None:
    project_dir = Path(args.project_dir)
    backend = open_backend(project_dir, args.backend, args.embedding_model, args.embedding_dims, require_existing=True)

    if args.keyword:
        results = backend.keyword_search(args.query, args.top_k)
    else:
        embedder = Embedder(args.embedding_model)
        query_embedding = embedder.embed([args.query])[0]
        results = backend.search(query_embedding, args.top_k)

    if hasattr(backend, "con"):
        backend.con.close()

    print(json.dumps({"results": results}))


def main() -> None:
    parser = argparse.ArgumentParser(description="Zolam v3 flat-file ingest/update/query pipeline")
    parser.add_argument("--mode", required=True, choices=["ingest", "update", "query"])
    parser.add_argument("--project-dir", required=True)
    parser.add_argument("--backend", required=True, choices=["duckdb", "jsonl"])
    parser.add_argument("--embedding-model", default=DEFAULT_MODEL)
    parser.add_argument("--embedding-dims", type=int, default=DEFAULT_DIMS)
    parser.add_argument("--files", nargs="*", default=[])
    parser.add_argument("--removed", nargs="*", default=[])
    parser.add_argument("--query")
    parser.add_argument("--top-k", type=int, default=5)
    parser.add_argument("--keyword", action="store_true")
    parser.add_argument("--output-json", action="store_true")
    args = parser.parse_args()

    try:
        if args.mode in ("ingest", "update"):
            run_ingest(args)
        else:
            if not args.query:
                raise RuntimeError("--query is required for --mode query")
            run_query(args)
    except Exception as e:  # noqa: BLE001 - top-level error boundary
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
